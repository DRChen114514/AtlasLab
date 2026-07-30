"""统一文档解析器 —— 多格式文档入口。

支持 PDF (MinerU API / 本地回退)、DOCX、HTML、Markdown、LaTeX。
为每个格式提供统一的 ParsedDocument 输出，包含结构化内容、元数据、缓存键。
"""
from __future__ import annotations

import hashlib
import json
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".html", ".md", ".txt", ".tex", ".latex"}


@dataclass
class ParsedDocument:
    """统一文档解析结果。"""
    source_path: str
    source_type: str  # "competition_problem", "reference_paper", "textbook", "dataset"
    text: str
    markdown: str = ""
    metadata: dict = field(default_factory=dict)
    images: list[bytes] = field(default_factory=list)
    tables: list[dict] = field(default_factory=list)
    formulas: list[str] = field(default_factory=list)
    sections: list[dict] = field(default_factory=list)
    chunks: list[str] = field(default_factory=list)
    cache_key: str = ""
    parsed_at: str = ""
    error: str = ""

    def __post_init__(self):
        if not self.parsed_at:
            self.parsed_at = datetime.now(timezone.utc).isoformat()


class DocumentParser:
    """统一文档解析器，根据文件后缀自动选择解析策略。

    用法:
        parser = DocumentParser()
        doc = parser.parse("problem.pdf")
        print(doc.text[:500])
    """

    def parse(self, file_path: str | Path, *, source_type: str = "unknown") -> ParsedDocument:
        """解析单个文档。

        Args:
            file_path: 文档路径。
            source_type: 文档类型标签。

        Returns:
            ParsedDocument 包含全文及结构化信息。
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix not in SUPPORTED_SUFFIXES:
            return ParsedDocument(
                source_path=str(file_path),
                source_type=source_type,
                text="",
                error=f"不支持的文件格式: {suffix}",
            )

        try:
            if suffix == ".pdf":
                return self._parse_pdf(file_path, source_type)
            elif suffix == ".docx":
                return self._parse_docx(file_path, source_type)
            elif suffix in (".html", ".htm"):
                return self._parse_html(file_path, source_type)
            elif suffix in (".md", ".txt"):
                return self._parse_text(file_path, source_type)
            elif suffix in (".tex", ".latex"):
                return self._parse_latex(file_path, source_type)
        except Exception as e:
            return ParsedDocument(
                source_path=str(file_path),
                source_type=source_type,
                text="",
                error=f"解析异常: {e}",
            )

        return ParsedDocument(
            source_path=str(file_path),
            source_type=source_type,
            text="",
            error=f"未处理的格式: {suffix}",
        )

    def _parse_pdf(self, path: Path, source_type: str) -> ParsedDocument:
        """PDF 解析：优先 MinerU，回退到本地 PyMuPDF。"""
        from math_agent.document.mineru import parse_via_mineru
        from math_agent.rag.ingest import _extract_pdf_text

        # 尝试 MinerU
        mineru_result = parse_via_mineru(path)
        if mineru_result.success:
            return ParsedDocument(
                source_path=str(path),
                source_type=source_type,
                text=mineru_result.markdown,
                markdown=mineru_result.markdown,
                metadata=mineru_result.metadata,
                cache_key=hashlib.sha256(mineru_result.markdown.encode()).hexdigest()[:16],
            )

        # 本地回退
        text = _extract_pdf_text(path)
        return ParsedDocument(
            source_path=str(path),
            source_type=source_type,
            text=text,
            metadata={"parser": "pymupdf_fallback", "mineru_error": mineru_result.error_message},
        )

    @staticmethod
    def _parse_docx(path: Path, source_type: str) -> ParsedDocument:
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            return ParsedDocument(
                source_path=str(path),
                source_type=source_type,
                text=text,
                metadata={"parser": "python-docx", "paragraph_count": len(paragraphs)},
            )
        except ImportError:
            return ParsedDocument(
                source_path=str(path),
                source_type=source_type,
                text="",
                error="python-docx 未安装",
            )

    @staticmethod
    def _parse_html(path: Path, source_type: str) -> ParsedDocument:
        try:
            from html.parser import HTMLParser

            class TextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text: list[str] = []
                    self._skip = False

                def handle_starttag(self, tag, attrs):
                    if tag in ("script", "style"):
                        self._skip = True

                def handle_endtag(self, tag):
                    if tag in ("script", "style"):
                        self._skip = False
                    if tag in ("p", "div", "br", "h1", "h2", "h3", "h4", "h5", "h6", "li"):
                        self.text.append("\n")

                def handle_data(self, data):
                    if not self._skip and data.strip():
                        self.text.append(data.strip())

            extractor = TextExtractor()
            extractor.feed(path.read_text(encoding="utf-8", errors="ignore"))
            return ParsedDocument(
                source_path=str(path),
                source_type=source_type,
                text="\n".join(extractor.text),
                metadata={"parser": "html.parser"},
            )
        except Exception as e:
            return ParsedDocument(
                source_path=str(path),
                source_type=source_type,
                text="",
                error=f"HTML 解析失败: {e}",
            )

    @staticmethod
    def _parse_text(path: Path, source_type: str) -> ParsedDocument:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return ParsedDocument(
            source_path=str(path),
            source_type=source_type,
            text=text,
            metadata={"parser": "plain_text"},
            cache_key=hashlib.sha256(text.encode()).hexdigest()[:16],
        )

    @staticmethod
    def _parse_latex(path: Path, source_type: str) -> ParsedDocument:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return ParsedDocument(
            source_path=str(path),
            source_type=source_type,
            text=text,
            metadata={"parser": "latex_raw"},
            cache_key=hashlib.sha256(text.encode()).hexdigest()[:16],
        )


def parse_document(
    file_path: str | Path,
    *,
    source_type: str = "unknown",
) -> ParsedDocument:
    """便捷函数：解析单个文档。"""
    return DocumentParser().parse(file_path, source_type=source_type)
